import streamlit as st
from docx import Document
from docx.shared import Inches
import io
import json

st.set_page_config(page_title="Firma de Consultoría - Bitácora Hito 1", layout="wide")

st.title("🏢 Firma de Consultoría en Ingeniería Industrial")
st.subheader("Plataforma de Bitácoras Web | Plantilla N° 01: Gate 0")

# ==============================================================================
# GESTIÓN DE RESPALDO Y RECUPERACIÓN DE AVANCE
# ==============================================================================
st.sidebar.header("💾 Gestión de Avance")

backup_file = st.sidebar.file_uploader("📂 Restaurar avance (subir .json)", type=["json"])

if backup_file is not None:
    try:
        data_loaded = json.load(backup_file)
        for key, val in data_loaded.items():
            st.session_state[key] = val
        st.sidebar.success("✅ Avance restaurado con éxito")
    except Exception as e:
        st.sidebar.error("❌ Error al cargar la copia de seguridad.")

# ==============================================================================
# BLOQUE 1 Y 2: HEADER Y ENTREGABLES
# ==============================================================================
st.markdown("### 📌 Bloque 1: Header de Control e Integrantes")
c1, c2 = st.columns(2)
with c1:
    integrante_1 = st.text_input("Integrante 1 (Director de Proyectos / General)", key="integrante_1")
    email_1 = st.text_input("Correo Electrónico Integrante 1", key="email_1")
    
    integrante_2 = st.text_input("Integrante 2 (Director de Mercado)", key="integrante_2")
    email_2 = st.text_input("Correo Electrónico Integrante 2", key="email_2")
    
    grupo = st.text_input("Grupo / Paralelo", key="grupo")

with c2:
    integrante_3 = st.text_input("Integrante 3 (Director Financiero y Riesgos)", key="integrante_3")
    email_3 = st.text_input("Correo Electrónico Integrante 3", key="email_3")
    
    integrante_4 = st.text_input("Integrante 4 (Director Técnico y Operaciones)", key="integrante_4")
    email_4 = st.text_input("Correo Electrónico Integrante 4", key="email_4")
    
    fecha_modulo = st.text_input("Fecha / Módulo", value="Mercado y Perfil del Proyecto", key="fecha_modulo")

st.markdown("### 🎯 Bloque 2: Matriz de Entregables e Hitos")
e1_col, e2_col = st.columns(2)
with e1_col:
    est_hito1 = st.selectbox("Estado Hito 1.1 (Selección Multicriterio)", ["Pendiente", "En Proceso", "Finalizado"], key="est_hito1")
    com_hito1 = st.text_input("Comentario Hito 1.1", key="com_hito1")
    
    est_hito2 = st.selectbox("Estado Hito 1.2 (Diagnóstico Causal y Alcance)", ["Pendiente", "En Proceso", "Finalizado"], key="est_hito2")
    com_hito2 = st.text_input("Comentario Hito 1.2", key="com_hito2")
with e2_col:
    est_hito3 = st.selectbox("Estado Hito 1.3 (Factibilidad y Riesgos)", ["Pendiente", "En Proceso", "Finalizado"], key="est_hito3")
    com_hito3 = st.text_input("Comentario Hito 1.3", key="com_hito3")

# ==============================================================================
# ROLES EN PESTAÑAS
# ==============================================================================
tab_mercado, tab_tecnico, tab_financiero, tab_general = st.tabs([
    "🟢 Director de Mercado", 
    "🟡 Director Técnico", 
    "🔴 Director Financiero y Riesgos", 
    "🎯 Director General"
])

# ------------------------------------------------------------------------------
# 🟢 DIRECTOR DE MERCADO
# ------------------------------------------------------------------------------
with tab_mercado:
    st.header("🟢 Director de Mercado")
    
    st.subheader("1. Sustento Teórico y Aplicación Técnica")
    conceptos_mercado = ["Diagnóstico causal", "Alcance técnico", "Problema", "Causa", "Efecto", "Árbol de problemas", "Storytelling"]
    
    for c in conceptos_mercado:
        st.markdown(f"**Concepto: {c}**")
        col_t1, col_t2, col_t3, col_t4 = st.columns(4)
        with col_t1:
            st.text_input(f"Definición - {c}", key=f"def_{c}")
        with col_t2:
            st.text_input(f"¿Cómo se usa en la práctica? - {c}", key=f"uso_{c}")
        with col_t3:
            st.text_input(f"¿Para qué sirve en este caso? - {c}", key=f"para_{c}")
        with col_t4:
            st.text_input(f"Fuente - {c}", key=f"src_{c}")

    st.subheader("2. Árbol de Problemas")
    prob_central = st.text_area("Descripción del Problema Central", key="prob_central")
    
    st.markdown("🖼️ **Evidencia Visual del Árbol de Problemas**")
    link_arbol = st.text_input("Enlace al Diagrama del Árbol (Miro / Lucidchart / Canva / Figma)", key="link_arbol")
    img_arbol = st.file_uploader("Cargar pantallazo / imagen del Árbol de Problemas", type=["png", "jpg", "jpeg"], key="img_arbol")
    
    st.markdown("**Efectos**")
    for i in range(1, 4):
        st.write(f"*Efecto {i}*")
        col_e1, col_e2, col_e3 = st.columns(3)
        with col_e1:
            st.text_input(f"Descripción Efecto {i}", key=f"desc_efecto_{i}")
        with col_e2:
            st.text_input(f"Evidencia y dato clave Efecto {i}", key=f"evid_efecto_{i}")
        with col_e3:
            st.text_input(f"Fuente Efecto {i}", key=f"src_efecto_{i}")

    st.markdown("**Causas Raíz**")
    for i in range(1, 4):
        st.write(f"*Causa Raíz {i}*")
        col_c1, col_c2, col_c3 = st.columns(3)
        with col_c1:
            st.text_input(f"Descripción Causa Raíz {i}", key=f"desc_causa_{i}")
        with col_c2:
            st.text_input(f"Evidencia y dato clave Causa {i}", key=f"evid_causa_{i}")
        with col_c3:
            st.text_input(f"Fuente Causa Raíz {i}", key=f"src_causa_{i}")

    st.subheader("3. Alcance Técnico")
    col_a1, col_a2 = st.columns(2)
    with col_a1:
        st.markdown("**Qué abarca el proyecto**")
        si_desc = st.text_area("Descripción (Qué abarca)", key="si_desc")
        si_just = st.text_area("Justificación (Qué abarca)", key="si_just")
    with col_a2:
        st.markdown("**Qué NO abarca el proyecto**")
        no_desc = st.text_area("Descripción (Qué NO abarca)", key="no_desc")
        no_just = st.text_area("Justificación (Qué NO abarca)", key="no_just")

    st.subheader("4. Storytelling (Guión)")
    storytelling_txt = st.text_area("Guión del Storytelling / Necesidad del Cliente", key="storytelling_txt")

# ------------------------------------------------------------------------------
# 🟡 DIRECTOR TÉCNICO Y OPERACIONES
# ------------------------------------------------------------------------------
with tab_tecnico:
    st.header("🟡 Director Técnico y Operaciones")
    
    st.subheader("1. Fundamentación Teórica y Aplicación Técnica")
    conceptos_tecnico = ["Matriz de pesos Ponderados", "Matriz AHP", "Criterios de Selección", "Alternativas de Proyecto"]
    
    for c in conceptos_tecnico:
        st.markdown(f"**Concepto: {c}**")
        col_t1, col_t2, col_t3, col_t4 = st.columns(4)
        with col_t1:
            st.text_input(f"Definición - {c}", key=f"def_tec_{c}")
        with col_t2:
            st.text_input(f"¿Cómo se usa en la práctica? - {c}", key=f"uso_tec_{c}")
        with col_t3:
            st.text_input(f"¿Para qué sirve en este caso? - {c}", key=f"para_tec_{c}")
        with col_t4:
            st.text_input(f"Fuente - {c}", key=f"src_tec_{c}")

    st.subheader("2. Definición de las 3 Propuestas de Proyecto Candidatas")
    st.markdown("Describa brevemente cada una de las 3 ideas de proyecto a evaluar:")
    propuesta_1 = st.text_input("Propuesta 1 (Idea A)", key="propuesta_1")
    propuesta_2 = st.text_input("Propuesta 2 (Idea B)", key="propuesta_2")
    propuesta_3 = st.text_input("Propuesta 3 (Idea C)", key="propuesta_3")

    st.subheader("3. Criterios de Evaluación y Matriz AHP")
    criterios_val = st.text_area("Criterios de Evaluación (Nombre, Peso % y Justificación)", key="criterios_val", 
                                 help="Ejemplo: Viabilidad Técnica (35%), Acceso a Datos (35%), Impacto Social/Económico (30%)")
    
    st.subheader("4. AI Log & Auditoría de Código")
    prompt_ia = st.text_area("Prompt utilizado para la IA / Script de Python", key="prompt_ia")
    audit_ia = st.text_area("Auditoría del Ingeniero: ¿Qué error, alucinación o supuesto irreal detectó en la respuesta de la IA y cómo lo corrigió?", key="audit_ia")
    
    st.subheader("5. Selección y Evaluación en Google Colab")
    link_colab = st.text_input("Enlace al Notebook de Colab Ejecutado", key="link_colab")
    img_colab = st.file_uploader("Cargar pantallazo de la ejecución en Colab / Gráfico de Resultados AHP", type=["png", "jpg", "jpeg"], key="img_colab")
    
    opcion_ganadora = st.text_area("🏆 Proyecto Ganador Seleccionado y Justificación Técnica", key="opcion_ganadora")

# ------------------------------------------------------------------------------
# 🔴 DIRECTOR FINANCIERO Y RIESGOS
# ------------------------------------------------------------------------------
with tab_financiero:
    st.header("🔴 Director Financiero y Riesgos")
    
    st.subheader("1. Fundamentación Teórica (Repositorio / Data Room)")
    col_f1, col_f2, col_f3, col_f4 = st.columns(4)
    with col_f1:
        def_repo = st.text_input("Definición (Repositorio)", key="def_repo")
    with col_f2:
        uso_repo = st.text_input("¿Cómo se usa en la práctica?", key="uso_repo")
    with col_f3:
        para_repo = st.text_input("¿Para qué sirve en este caso?", key="para_repo")
    with col_f4:
        fuent_repo = st.text_input("Fuente (Repositorio)", key="fuent_repo")
        
    link_repositorio = st.text_input("Link al Repositorio / Data Room (Google Drive / Sheets)", key="link_repositorio")
    img_dataroom = st.file_uploader("Cargar captura de pantalla del Data Room", type=["png", "jpg", "jpeg"], key="img_data")

    st.subheader("2. Validación de Datos Reales")
    val_items = ["Datos de mercado y precios", "Datos técnicos operativos", "Estructura de Costos", "Funcionamiento Matriz"]
    
    for item in val_items:
        st.markdown(f"**{item}**")
        col_v1, col_v2, col_v3 = st.columns(3)
        with col_v1:
            st.selectbox(f"¿Existe? - {item}", ["Sí", "No"], key=f"ex_{item}")
        with col_v2:
            st.text_input(f"Evidencia de fuentes - {item}", key=f"ev_{item}")
        with col_v3:
            st.text_input(f"Fuente - {item}", key=f"fu_{item}")

# ------------------------------------------------------------------------------
# 🎯 DIRECTOR GENERAL (INTEGRACIÓN, GUARDADO Y DESCARGA)
# ------------------------------------------------------------------------------
with tab_general:
    st.header("🎯 Director General")
    
    nombre_formal = st.text_input("Nombre Formal del Proyecto Ganador", key="nombre_formal")
    resumen_proyecto = st.text_area("Resumen del Proyecto", key="resumen_proyecto")
    
    st.markdown("**Descripción del Cliente**")
    col_g1, col_g2 = st.columns(2)
    with col_g1:
        desc_cliente = st.text_area("Descripción del Cliente", key="desc_cliente")
    with col_g2:
        fuent_cliente = st.text_input("Fuente", key="fuent_cliente")
        evid_cliente = st.text_input("Evidencia", key="evid_cliente")
        
    objetivo_general = st.text_area("Objetivo General Aprobado", key="objetivo_general")
    
    st.markdown("---")
    st.subheader("💾 Guardar Avance / Restauración")

    backup_data = {k: v for k, v in st.session_state.items() if isinstance(v, (str, int, float, list, dict))}
    json_bytes = json.dumps(backup_data, ensure_ascii=False, indent=2).encode('utf-8')

    st.download_button(
        label="📥 Descargar Avance (Copia de Seguridad .json)",
        data=json_bytes,
        file_name="Avance_Bitacora_Hito_1.json",
        mime="application/json"
    )

    st.markdown("---")
    # ==========================================================================
    # GENERADOR DE WORD (.DOCX) - INFORME COMPLETO
    # ==========================================================================
    if st.button("🚀 Generar Bitácora Completa (.docx)"):
        doc = Document()
        
        # Estilos y Encabezado
        doc.add_heading('FIRMA DE CONSULTORÍA EN INGENIERÍA INDUSTRIAL', 0)
        doc.add_heading('Bitácora N° 01: Identificación de Oportunidades (Gate 0)', level=1)
        
        # ---------------------------------------------------------
        # BLOQUE 1: RESUMEN EJECUTIVO Y DATOS
        # ---------------------------------------------------------
        doc.add_heading('1. Resumen Ejecutivo', level=2)
        doc.add_paragraph(f"Proyecto Ganador: {st.session_state.get('nombre_formal', 'No definido')}").bold = True
        doc.add_paragraph(f"Objetivo General: {st.session_state.get('objetivo_general', '')}")
        doc.add_paragraph(f"Resumen: {st.session_state.get('resumen_proyecto', '')}")
        
        doc.add_paragraph("Perfil del Cliente:").bold = True
        doc.add_paragraph(f"{st.session_state.get('desc_cliente', '')} (Evidencia: {st.session_state.get('evid_cliente', '')} | Fuente: {st.session_state.get('fuent_cliente', '')})")

        doc.add_heading('1.1. Datos del Equipo', level=3)
        p_equipo = doc.add_paragraph()
        p_equipo.add_run(f"Grupo / Paralelo: {st.session_state.get('grupo', 'No definido')}\n").bold = True
        p_equipo.add_run(f"Director General: {st.session_state.get('integrante_1', '')} - {st.session_state.get('email_1', '')}\n")
        p_equipo.add_run(f"Director de Mercado: {st.session_state.get('integrante_2', '')} - {st.session_state.get('email_2', '')}\n")
        p_equipo.add_run(f"Director Financiero: {st.session_state.get('integrante_3', '')} - {st.session_state.get('email_3', '')}\n")
        p_equipo.add_run(f"Director Técnico: {st.session_state.get('integrante_4', '')} - {st.session_state.get('email_4', '')}\n")
        p_equipo.add_run(f"Fecha: {st.session_state.get('fecha_modulo', '')}")

        doc.add_heading('1.2. Estado de Hitos', level=3)
        doc.add_paragraph(f"Hito 1.1 (Selección Multicriterio): {st.session_state.get('est_hito1', '')} - Comentario: {st.session_state.get('com_hito1', '')}", style='List Bullet')
        doc.add_paragraph(f"Hito 1.2 (Diagnóstico y Alcance): {st.session_state.get('est_hito2', '')} - Comentario: {st.session_state.get('com_hito2', '')}", style='List Bullet')
        doc.add_paragraph(f"Hito 1.3 (Factibilidad y Riesgos): {st.session_state.get('est_hito3', '')} - Comentario: {st.session_state.get('com_hito3', '')}", style='List Bullet')

        # ---------------------------------------------------------
        # MARCO TEÓRICO UNIFICADO
        # ---------------------------------------------------------
        doc.add_heading('2. Marco Teórico y Aplicación Técnica', level=2)
        doc.add_paragraph("Sustento teórico de las herramientas utilizadas (Mercado, Técnico y Financiero) y su justificación de aplicación al proyecto.")
        
        # Consolidar teoría incluyendo la del Financiero
        conceptos_totales = [
            ("Mercado", conceptos_mercado, "def_", "uso_", "para_", "src_"),
            ("Técnico", conceptos_tecnico, "def_tec_", "uso_tec_", "para_tec_", "src_tec_"),
            ("Financiero", ["Repositorio / Data Room"], "def_repo", "uso_repo", "para_repo", "fuent_repo")
        ]
        
        table_teoria = doc.add_table(rows=1, cols=5)
        table_teoria.style = 'Table Grid'
        hdr_cells = table_teoria.rows[0].cells
        hdr_cells[0].text = 'Concepto'
        hdr_cells[1].text = 'Definición'
        hdr_cells[2].text = '¿Cómo se usa?'
        hdr_cells[3].text = '¿Para qué sirve?'
        hdr_cells[4].text = 'Fuente'
        
        for area, conceptos, pre_def, pre_uso, pre_para, pre_src in conceptos_totales:
            for c in conceptos:
                row_cells = table_teoria.add_row().cells
                row_cells[0].text = c
                
                # Manejo especial para Financiero que no usa sufijo en el iterador
                if area == "Financiero":
                    row_cells[1].text = st.session_state.get("def_repo", "")
                    row_cells[2].text = st.session_state.get("uso_repo", "")
                    row_cells[3].text = st.session_state.get("para_repo", "")
                    row_cells[4].text = st.session_state.get("fuent_repo", "")
                else:
                    row_cells[1].text = st.session_state.get(f"{pre_def}{c}", "")
                    row_cells[2].text = st.session_state.get(f"{pre_uso}{c}", "")
                    row_cells[3].text = st.session_state.get(f"{pre_para}{c}", "")
                    row_cells[4].text = st.session_state.get(f"{pre_src}{c}", "")

        # ---------------------------------------------------------
        # DIRECTOR DE MERCADO (Árbol y Alcance)
        # ---------------------------------------------------------
        doc.add_heading('3. Diagnóstico de Mercado', level=2)
        
        doc.add_paragraph(f"Problema Central: {st.session_state.get('prob_central', '')}")
        
        doc.add_paragraph("En la Figura 1 se observa el diagrama del Árbol de Problemas estructurado por la firma de consultoría.")
        
        img_arbol = st.session_state.get('img_arbol')
        if img_arbol is not None:
            image_stream = io.BytesIO(img_arbol.getvalue())
            doc.add_picture(image_stream, width=Inches(6.0))
            fig1 = doc.add_paragraph("Figura 1. Diagrama del Árbol de Problemas.")
            fig1.alignment = 1 
            fig1.runs[0].font.italic = True
            
        link_arbol = st.session_state.get('link_arbol', '')
        if link_arbol:
            doc.add_paragraph(f"Enlace de trabajo interactivo (Miro/Lucidchart): {link_arbol}")

        doc.add_paragraph("Causas del problema detectado:")
        for i in range(1, 4):
            desc_c = st.session_state.get(f"desc_causa_{i}", "")
            if desc_c:
                doc.add_paragraph(f"Causa {i}: {desc_c}. (Evidencia: {st.session_state.get(f'evid_causa_{i}', '')} | Fuente: {st.session_state.get(f'src_causa_{i}', '')})", style='List Bullet')

        doc.add_paragraph("Efectos principales en el entorno:")
        for i in range(1, 4):
            desc_e = st.session_state.get(f"desc_efecto_{i}", "")
            if desc_e:
                doc.add_paragraph(f"Efecto {i}: {desc_e}. (Evidencia: {st.session_state.get(f'evid_efecto_{i}', '')} | Fuente: {st.session_state.get(f'src_efecto_{i}', '')})", style='List Bullet')

        doc.add_heading('3.1. Alcance Técnico y Storytelling', level=3)
        doc.add_paragraph(f"El proyecto SÍ abarca: {st.session_state.get('si_desc', '')}. Justificación: {st.session_state.get('si_just', '')}", style='List Bullet')
        doc.add_paragraph(f"El proyecto NO abarca: {st.session_state.get('no_desc', '')}. Justificación: {st.session_state.get('no_just', '')}", style='List Bullet')
        
        doc.add_paragraph("Guión de Storytelling (Necesidad del Cliente):").bold = True
        doc.add_paragraph(st.session_state.get('storytelling_txt', ''), style='Intense Quote')

        # ---------------------------------------------------------
        # DIRECTOR TÉCNICO (AHP e IA)
        # ---------------------------------------------------------
        doc.add_heading('4. Evaluación Técnica (AHP) y Auditoría de IA', level=2)
        
        doc.add_paragraph("Las alternativas de proyecto evaluadas fueron:")
        doc.add_paragraph(f"Propuesta 1: {st.session_state.get('propuesta_1', '')}", style='List Number')
        doc.add_paragraph(f"Propuesta 2: {st.session_state.get('propuesta_2', '')}", style='List Number')
        doc.add_paragraph(f"Propuesta 3: {st.session_state.get('propuesta_3', '')}", style='List Number')

        doc.add_paragraph(f"Criterios de Evaluación y Pesos: {st.session_state.get('criterios_val', '')}")
        
        doc.add_paragraph("En la Figura 2 se presenta la ejecución del algoritmo AHP.")
        img_colab = st.session_state.get('img_colab')
        if img_colab is not None:
            image_stream2 = io.BytesIO(img_colab.getvalue())
            doc.add_picture(image_stream2, width=Inches(6.0))
            fig2 = doc.add_paragraph("Figura 2. Ejecución y Resultados del Código AHP (Colab).")
            fig2.alignment = 1
            fig2.runs[0].font.italic = True
            
        link_colab = st.session_state.get('link_colab', '')
        if link_colab:
            doc.add_paragraph(f"Enlace al Notebook de Colab: {link_colab}")

        doc.add_heading('4.1 Log de Inteligencia Artificial', level=3)
        doc.add_paragraph(f"Prompt utilizado: {st.session_state.get('prompt_ia', '')}")
        doc.add_paragraph(f"Auditoría del Ingeniero: {st.session_state.get('audit_ia', '')}", style='Intense Quote')

        # ---------------------------------------------------------
        # DIRECTOR FINANCIERO (Data Room y Validación)
        # ---------------------------------------------------------
        doc.add_heading('5. Estructura Financiera y Gestión de Datos', level=2)
        
        doc.add_paragraph("Para garantizar la trazabilidad de la información, se estructuró un Repositorio / Data Room. (Ver Figura 3).")
        img_dataroom = st.session_state.get('img_data')
        if img_dataroom is not None:
            image_stream3 = io.BytesIO(img_dataroom.getvalue())
            doc.add_picture(image_stream3, width=Inches(6.0))
            fig3 = doc.add_paragraph("Figura 3. Estructura del Data Room Financiero y Técnico.")
            fig3.alignment = 1
            fig3.runs[0].font.italic = True

        link_repo = st.session_state.get('link_repositorio', '')
        if link_repo:
            doc.add_paragraph(f"Enlace al Repositorio (Google Drive/Sheets): {link_repo}")

        doc.add_heading('5.1 Validación de Datos Reales', level=3)
        
        table_val = doc.add_table(rows=1, cols=4)
        table_val.style = 'Table Grid'
        hv_cells = table_val.rows[0].cells
        hv_cells[0].text = 'Item de Validación'
        hv_cells[1].text = '¿Existe?'
        hv_cells[2].text = 'Evidencia'
        hv_cells[3].text = 'Fuente'

        val_items = ["Datos de mercado y precios", "Datos técnicos operativos", "Estructura de Costos", "Funcionamiento Matriz"]
        for item in val_items:
            rv_cells = table_val.add_row().cells
            rv_cells[0].text = item
            rv_cells[1].text = st.session_state.get(f"ex_{item}", "")
            rv_cells[2].text = st.session_state.get(f"ev_{item}", "")
            rv_cells[3].text = st.session_state.get(f"fu_{item}", "")

        # Descarga final
        bio = io.BytesIO()
        doc.save(bio)
        
        st.success("¡Bitácora generada con éxito en formato de informe técnico completo!")
        st.download_button(
            label="📄 Descargar Informe Técnico (.docx)",
            data=bio.getvalue(),
            file_name=f"Informe_Tecnico_Hito_1_{st.session_state.get('grupo', 'Proyecto')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
